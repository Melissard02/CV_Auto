from docx import Document
import os
import glob
from docx2pdf import convert


def replace_placeholder_text(doc, replacements):
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for placeholder, new_text in replacements.items():
            if placeholder in text:
                text = text.replace(placeholder, new_text)
        # Rebuild paragraph so replacements actually show up
        for i in range(len(paragraph.runs)):
            paragraph.runs[i].text = ""
        paragraph.add_run(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_text(cell, replacements)


def get_inputs():
    jobs = []
    print("Hi! Please input job information like this: ")
    print("Example:\nCompany: Apple\nJob Title: Software Engineer\n")

    while True:
        company = input("Company Name: ")
        job_title = input("Job Title: ")

        confirm = input(f"\nIs this correct?\nCompany: {company}\nJob Title: {job_title}\n(y/n)\n> ")
        if confirm.lower() == "y":
            jobs.append((company, job_title))
            print(f"\nAdded: {company} — {job_title}\n")
        else:
            print("Please try again!\n")
            continue

        again = input("Add another? (y/n)\n> ")
        if again.lower() == "n":
            break
    return jobs


def generate_cover_letters(jobs):
    os.makedirs("generated_letters", exist_ok=True)

    for company, job_title in jobs:
        doc = Document("cover_letter_template.docx")

        replacements = {
            "[COMPANY_NAME]": company,
            "[JOB_TITLE]": job_title
        }

        replace_placeholder_text(doc, replacements)

        filename = f"{company}_{job_title}.docx".replace(" ", "_")
        filepath = os.path.join("generated_letters", filename)
        doc.save(filepath)
        print(f"Created: {filename}")

    print("\nAll DOCX letters done!")


def convert_to_pdfs():
    print("\nConverting new DOCXs to PDFs...")
    docx_files = glob.glob("generated_letters/*.docx")
    for docx_file in docx_files:
        pdf_file = docx_file.replace(".docx", ".pdf")
        if not os.path.exists(pdf_file):
            convert(docx_file, os.path.dirname(docx_file))
            print(f"Converted: {os.path.basename(pdf_file)}")
    print("PDFs are ready! Check your generated_letters folder")


def main():
    # Cross-platform console clear
    os.system('cls' if os.name == 'nt' else 'clear')

    jobs = get_inputs()
    generate_cover_letters(jobs)
    convert_to_pdfs()


if __name__ == "__main__":
    main()
